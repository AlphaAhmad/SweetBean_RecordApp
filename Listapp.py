import sys
import qrcode
from PyQt5.QtWidgets import *
from PyQt5.QtGui import QPixmap, QImage, QFont, QTextOption
from PyQt5.QtCore import Qt, QTimer, QDateTime, QSettings
from docx import Document
from docx.shared import Inches

# Global variables
headers = ["Product", "Quantity (Kg)", "Price (Rs)", "Actions"]
Policy1 = "Purchase sample before buying product. Bought product won't be refunded or Exchanged."
Policy2 = "Only transfer online payment to the mentioned account number and payment detail."
Fixed_Policy_List = [Policy1,Policy2]
Whatsapp_link = "https://wa.me/message/A3RKKMMSAZGIF1"
Insta_link = "https://www.instagram.com/sweetbean.info?utm_source=qr&igsh=cTJmY3IwaHlwNzJq"

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Sweet Bean")
        self.setGeometry(100, 100, 700, 700)  # Optimal window size

        # Initialize settings and last invoice number
        self.settings = QSettings("Sweet Bean", "InvoiceApp")
        self.last_invoice_number = self.settings.value("last_invoice_number", 0, int)
        if self.last_invoice_number == 0:
            self.last_invoice_number = 1

        # Fixed policies that should always be loaded but not saved
        # self.fixed_policies = [
        #     "No refunds after 7 days of purchase.",
        #     "Customers must retain the receipt for exchanges."
        # ]

        # Load user-defined policies from settings (excluding fixed ones)
        saved_policies = self.settings.value("policies", [], list)
        
        # Filter out duplicates (Avoid saving fixed policies)
        self.policies = [p for p in saved_policies if p not in Fixed_Policy_List]

        # Set up the UI
        self.setup_ui()

        # Start the timer for auto-updating date and time
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_time)
        self.timer.start(1000)


    def setup_ui(self):
        """Set up the main UI components."""
        
        # Main content widget
        self.main_widget = QWidget()
        self.main_layout = QVBoxLayout(self.main_widget)
        self.main_layout.setContentsMargins(20, 20, 20, 40)
        self.main_layout.setSpacing(15)

        # Sections of the UI
        self.main_layout.addLayout(self.create_metadata_layout())   # Logo, Date, Invoice, NTN
        self.main_layout.addLayout(self.create_input_layout())      # Input fields + Add button

        # Table for product list
        self.table = self.create_table()
        self.main_layout.addWidget(self.table)

        # Total sum widget
        self.total_sum_widget = self.create_total_sum_widget()
        self.main_layout.addWidget(self.total_sum_widget)

        # Payment section
        self.main_layout.addLayout(self.create_payment_layout())       # Payment method & account
        self.main_layout.addLayout(self.create_payment_status_layout()) # Payment status radio buttons

        # QR codes
        self.main_layout.addLayout(self.create_qr_layout())

        # Policy Notes Section
        self.main_layout.addWidget(self.create_policy_layout())

        # Save button
        self.save_button = QPushButton("Save as DOC")
        self.save_button.clicked.connect(self.save_as_doc)
        self.main_layout.addWidget(self.save_button)

        # Scrollable Area for Large Content
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setWidget(self.main_widget)

        # Set the central widget
        self.setCentralWidget(self.scroll_area)
        # Apply dark theme stylesheet
        self.setStyleSheet("""
            QMainWindow {
                background-color: #2d2d2d;
            }
            QLabel {
                font-size: 14px;
                color: #ffffff;
            }
            QLineEdit, QSpinBox, QComboBox {
                font-size: 14px;
                padding: 8px;
                border: 1px solid #555;
                border-radius: 4px;
                background-color: #3d3d3d;
                color: #ffffff;
            }
            QPushButton {
                font-size: 14px;
                padding: 8px 16px;
                background-color: #007bff;
                color: white;
                border: none;
                border-radius: 4px;
            }
            QPushButton:hover {
                background-color: #0056b3;
            }
            QTableWidget {
                font-size: 14px;
                border: 1px solid #555;
                border-radius: 4px;
                background-color: #3d3d3d;
                color: #ffffff;
            }
            QTableWidget::item {
                padding: 8px;
            }
            QHeaderView::section {
                background-color: #007bff;
                color: white;
                padding: 8px;
                font-size: 14px;
            }
            QRadioButton {
                font-size: 14px;
                color: #ffffff;
            }
            QLineEdit[readOnly="true"] {
                background-color: #555;
                color: #ffffff;
            }
            QScrollArea {
                border: none;
            }
        """)

    def create_metadata_layout(self):
        """Create the metadata layout (logo, date/time, invoice, NTN, and Reset Button)."""
        layout = QGridLayout()
        layout.setSpacing(15)

        # Reset Invoice Button (Top Left)
        self.reset_invoice_btn = QPushButton("Reset Invoice")
        self.reset_invoice_btn.clicked.connect(self.reset_invoice_number)
        layout.addWidget(self.reset_invoice_btn, 0, 0, 1, 2)  # Positioned at the top left

        # Logo
        logo_label = QLabel()
        logo = QPixmap("logo.jpeg")
        scaled_logo = logo.scaled(120, 120, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        logo_label.setPixmap(scaled_logo)
        logo_label.setAlignment(Qt.AlignCenter)
        layout.addWidget(logo_label, 1, 0, 1, 2)

        # Date and Time
        self.date_time_label = QLabel()
        self.date_time_label.setAlignment(Qt.AlignCenter)
        self.date_time_label.setFont(QFont("Arial", 12))
        layout.addWidget(self.date_time_label, 2, 0, 1, 2)

        # Invoice Number (Separate Row)
        invoice_text_label = QLabel("Invoice Number:")
        invoice_text_label.setFont(QFont("Arial", 12, QFont.Bold))
        self.invoice_label = QLabel(f"{self.last_invoice_number}")
        self.invoice_label.setFont(QFont("Arial", 12, QFont.Bold))
        layout.addWidget(invoice_text_label, 3, 0)
        layout.addWidget(self.invoice_label, 3, 1)

        # NTN (Separate Row)
        self.ntn_text_label = QLineEdit() # TODO
        self.ntn_text_label.setPlaceholderText("Enter NTN number")
        self.ntn_text_label.setFont(QFont("Arial", 12, QFont.Bold))
        layout.addWidget(self.ntn_text_label, 4, 0)
        return layout


    def create_input_layout(self):
        """Create the input layout (product name, quantity, price, add button)."""
        layout = QHBoxLayout()
        layout.setSpacing(10)

        self.name_input = QLineEdit()
        self.name_input.setPlaceholderText("Product Name")

        self.quantity_input = QSpinBox()
        self.quantity_input.setRange(0, 100000000)
        self.quantity_input.setPrefix("Kg ")

        self.price_input = QSpinBox()
        self.price_input.setRange(0, 1000000000)
        self.price_input.setPrefix("Rs ")

        add_button = QPushButton("Add Product")
        add_button.clicked.connect(self.add_row)

        layout.addWidget(self.name_input)
        layout.addWidget(self.quantity_input)
        layout.addWidget(self.price_input)
        layout.addWidget(add_button)

        return layout

    def create_table(self):
        """Create and configure the table for product list."""
        table = QTableWidget()
        table.setColumnCount(4)
        table.setHorizontalHeaderLabels(headers)
        table.horizontalHeader().setStretchLastSection(True)
        table.verticalHeader().setVisible(False)
        table.setSelectionBehavior(QTableWidget.SelectRows)
        table.setEditTriggers(QTableWidget.NoEditTriggers)

        # Set row height to make rows bigger
        table.verticalHeader().setDefaultSectionSize(40)  # Adjust row height

        # Set table size to show at least 3-4 rows
        table.setMinimumHeight(160)  # 4 rows * 40px height

        return table

    def create_total_sum_widget(self):
        """Create the total sum widget."""
        total_sum_widget = QLineEdit()
        total_sum_widget.setPlaceholderText("Total Sum (Rs)")
        total_sum_widget.setReadOnly(True)
        total_sum_widget.setAlignment(Qt.AlignRight)
        total_sum_widget.setFont(QFont("Arial", 12, QFont.Bold))
        return total_sum_widget

    def create_payment_layout(self):
        """Create the payment method and account number layout."""
        layout = QVBoxLayout()
        layout.setSpacing(10)

        self.payment_method_input = QLineEdit()
        self.payment_method_input.setPlaceholderText("Enter Payment Method")

        self.account_number_input = QLineEdit()
        self.account_number_input.setPlaceholderText("Enter Account Number")

        layout.addWidget(QLabel("Payment Details"))
        layout.addWidget(self.payment_method_input)
        layout.addWidget(self.account_number_input)

        return layout

    def create_payment_status_layout(self):
        """Create the payment status radio buttons layout."""
        layout = QVBoxLayout()
        layout.setSpacing(10)

        self.paid_radio = QRadioButton("Paid")
        self.pending_radio = QRadioButton("Pending")

        button_group = QButtonGroup(self)
        button_group.addButton(self.paid_radio)
        button_group.addButton(self.pending_radio)

        layout.addWidget(QLabel("Payment Status"))
        layout.addWidget(self.paid_radio)
        layout.addWidget(self.pending_radio)

        return layout

    def create_qr_layout(self):
        """Create the QR codes layout."""
        layout = QVBoxLayout()
        layout.setSpacing(10)

        # Generate and display QR codes
        Whatsapp_qr = self.generate_qr(Whatsapp_link)
        instagram_qr = self.generate_qr(Insta_link)

        fb_qr_label = QLabel()
        fb_qr_label.setPixmap(self.qr_to_pixmap(Whatsapp_qr))
        fb_qr_label.setAlignment(Qt.AlignCenter)

        ig_qr_label = QLabel()
        ig_qr_label.setPixmap(self.qr_to_pixmap(instagram_qr))
        ig_qr_label.setAlignment(Qt.AlignCenter)

        # Labels for QR codes
        fb_label = QLabel("Visit our Whatsapp")
        ig_label = QLabel("Visit our Instagram")
        fb_label.setAlignment(Qt.AlignCenter)
        ig_label.setAlignment(Qt.AlignCenter)

        # Add QR codes and labels to the layout
        qr_images_layout = QHBoxLayout()
        qr_images_layout.addWidget(fb_qr_label)
        qr_images_layout.addWidget(ig_qr_label)

        qr_labels_layout = QHBoxLayout()
        qr_labels_layout.addWidget(fb_label)
        qr_labels_layout.addWidget(ig_label)

        layout.addWidget(QLabel("Scan QR Codes"))
        layout.addLayout(qr_images_layout)
        layout.addLayout(qr_labels_layout)

        return layout

    def create_policy_layout(self):
        """Create the policy notes layout."""
        layout = QVBoxLayout()
        layout.setSpacing(10)

        self.policy1_input = QLineEdit(Policy1)
        self.policy1_input.setEnabled(False)

        self.policy2_input = QLineEdit(Policy2)
        self.policy2_input.setEnabled(False)

        layout.addWidget(QLabel("Policies"))
        layout.addWidget(self.policy1_input)
        layout.addWidget(self.policy2_input)

        return layout

    def generate_qr(self, url):
        """Generate a QR code for the given URL."""
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=4,
            border=4,
        )
        qr.add_data(url)
        qr.make(fit=True)
        return qr.make_image(fill="black", back_color="white")

    def qr_to_pixmap(self, qr_img):
        """Convert a QR code image to a QPixmap."""
        qr_img = qr_img.convert("RGB")
        width, height = qr_img.size
        data = qr_img.tobytes("raw", "RGB")
        qimage = QImage(data, width, height, QImage.Format_RGB888)
        return QPixmap.fromImage(qimage)

    def add_row(self):
        """Add a new row to the table."""
        product_name = self.name_input.text()
        quantity = self.quantity_input.value()
        price = self.price_input.value()

        if not product_name:
            QMessageBox.warning(self, "Error", "Product Name field is empty.")
            return

        if not quantity:
            QMessageBox.warning(self, "Error", "Quantity field is empty.")
            return

        if not price:
            QMessageBox.warning(self, "Error", "Price field is empty.")
            return

        row_index = self.table.rowCount()
        self.table.insertRow(row_index)

        self.table.setItem(row_index, 0, QTableWidgetItem(product_name))
        self.table.setItem(row_index, 1, QTableWidgetItem(f"{quantity}"))
        self.table.setItem(row_index, 2, QTableWidgetItem(f"{price}"))

        # Create a smaller delete button
        delete_button = QPushButton("Delete")
        delete_button.setFixedSize(90, 25)  # Set button size (width, height)
        delete_button.clicked.connect(self.delete_row)

        # Center the text inside the button using stylesheet
        delete_button.setStyleSheet("""
            QPushButton {
                text-align: center;  /* Center the text horizontally */
                padding: 0px;        /* Remove padding to ensure text is centered */
                margin: 0px;         /* Remove margin to ensure text is centered */
            }
        """)

        # Center the button in its column
        button_widget = QWidget()
        button_layout = QHBoxLayout(button_widget)
        button_layout.addWidget(delete_button)
        button_layout.setAlignment(Qt.AlignCenter)  # Center the button
        button_layout.setContentsMargins(0, 0, 0, 0)  # Remove margins

        self.table.setCellWidget(row_index, 3, button_widget)

        # Clear input fields
        self.name_input.clear()
        self.quantity_input.setValue(0)
        self.price_input.setValue(0)

        # Update the total sum
        self.update_total_sum()

    def delete_row(self):
        """Delete a row from the table."""
        button = self.sender()
        if button:
            index = self.table.indexAt(button.pos())
            if index.isValid():
                self.table.removeRow(index.row())
                self.update_total_sum()

    def update_time(self):
        """Update the date and time label."""
        current_time = QDateTime.currentDateTime().toString("dd-MM-yyyy  hh:mm:ss")
        self.date_time_label.setText(current_time)

    def update_invoice_label(self):
        """Update the invoice label."""
        self.invoice_label.setText(f"Invoice number: {self.last_invoice_number}")

    def increment_Invoice_number(self):
        self.last_invoice_number+=1
        self.update_invoice_label()

    def reset_invoice_number(self):
        """Reset invoice number to 1 and update the UI."""
        self.last_invoice_number = 1
        self.settings.setValue("last_invoice_number", self.last_invoice_number)  # Store in QSettings
        self.invoice_label.setText(str(self.last_invoice_number))  # Update UI
        QMessageBox.information(self, "Invoice Reset", "Invoice number has been reset to 1.")


    def update_total_sum(self):
        """Calculate and update the total sum of prices in the table."""
        total_sum = 0
        for row in range(self.table.rowCount()):
            price_item = self.table.item(row, 2)
            if price_item:
                total_sum += int(price_item.text())
        self.total_sum_widget.setText(f"Total Sum: Rs {total_sum}")

    
    def closeEvent(self, event):
        """Save the last invoice number when the window is closed."""
        self.settings.setValue("last_invoice_number", self.last_invoice_number)
        event.accept()

    def create_table(self):
        """Create and configure the table for product list."""
        table = QTableWidget()
        table.setColumnCount(4)
        table.setHorizontalHeaderLabels(headers)
        table.horizontalHeader().setStretchLastSection(True)
        table.verticalHeader().setVisible(False)
        table.setSelectionBehavior(QTableWidget.SelectRows)
        table.setEditTriggers(QTableWidget.NoEditTriggers)

        # Set row height to make rows bigger
        table.verticalHeader().setDefaultSectionSize(40)  # Adjust row height

        # Set table size to show at least 3-4 rows
        table.setMinimumHeight(160)  # 4 rows * 40px height

        return table

    def create_total_sum_widget(self):
        """Create the total sum widget."""
        total_sum_widget = QLineEdit()
        total_sum_widget.setPlaceholderText("Total Sum (Rs)")
        total_sum_widget.setReadOnly(True)
        total_sum_widget.setAlignment(Qt.AlignRight)
        total_sum_widget.setFont(QFont("Arial", 12, QFont.Bold))
        return total_sum_widget

    def create_payment_layout(self):
        """Create the payment method and account number layout."""
        layout = QVBoxLayout()
        layout.setSpacing(10)

        self.payment_method_input = QLineEdit()
        self.payment_method_input.setPlaceholderText("Enter Payment Method")

        self.account_number_input = QLineEdit()
        self.account_number_input.setPlaceholderText("Enter Account Number")

        layout.addWidget(QLabel("Payment Details"))
        layout.addWidget(self.payment_method_input)
        layout.addWidget(self.account_number_input)

        return layout

    def create_payment_status_layout(self):
        """Create the payment status radio buttons layout."""
        layout = QVBoxLayout()
        layout.setSpacing(10)

        self.paid_radio = QRadioButton("Paid")
        self.pending_radio = QRadioButton("Pending")

        button_group = QButtonGroup(self)
        button_group.addButton(self.paid_radio)
        button_group.addButton(self.pending_radio)

        layout.addWidget(QLabel("Payment Status"))
        layout.addWidget(self.paid_radio)
        layout.addWidget(self.pending_radio)

        return layout

    def create_qr_layout(self):
        """Create the QR codes layout."""
        layout = QVBoxLayout()
        layout.setSpacing(10)

        # Generate and display QR codes
        Whatsapp_qr = self.generate_qr(Whatsapp_link)
        instagram_qr = self.generate_qr(Insta_link)

        fb_qr_label = QLabel()
        fb_qr_label.setPixmap(self.qr_to_pixmap(Whatsapp_qr))
        fb_qr_label.setAlignment(Qt.AlignCenter)

        ig_qr_label = QLabel()
        ig_qr_label.setPixmap(self.qr_to_pixmap(instagram_qr))
        ig_qr_label.setAlignment(Qt.AlignCenter)

        # Labels for QR codes
        fb_label = QLabel("Visit our Whatsapp")
        ig_label = QLabel("Visit our Instagram")
        fb_label.setAlignment(Qt.AlignCenter)
        ig_label.setAlignment(Qt.AlignCenter)

        # Add QR codes and labels to the layout
        qr_images_layout = QHBoxLayout()
        qr_images_layout.addWidget(fb_qr_label)
        qr_images_layout.addWidget(ig_qr_label)

        qr_labels_layout = QHBoxLayout()
        qr_labels_layout.addWidget(fb_label)
        qr_labels_layout.addWidget(ig_label)

        layout.addWidget(QLabel("Scan QR Codes"))
        layout.addLayout(qr_images_layout)
        layout.addLayout(qr_labels_layout)

        return layout

    # def create_policy_layout(self):
    #     """Create the policy notes layout."""
    #     layout = QVBoxLayout()
    #     layout.setSpacing(10)

    #     self.policy1_input = QLineEdit(Policy1)
    #     self.policy1_input.setEnabled(False)

    #     self.policy2_input = QLineEdit(Policy2)
    #     self.policy2_input.setEnabled(False)

    #     layout.addWidget(QLabel("Policies"))
    #     layout.addWidget(self.policy1_input)
    #     layout.addWidget(self.policy2_input)

    #     return layout

    def generate_qr(self, url):
        """Generate a QR code for the given URL."""
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=4,
            border=4,
        )
        qr.add_data(url)
        qr.make(fit=True)
        return qr.make_image(fill="black", back_color="white")

    def qr_to_pixmap(self, qr_img):
        """Convert a QR code image to a QPixmap."""
        qr_img = qr_img.convert("RGB")
        width, height = qr_img.size
        data = qr_img.tobytes("raw", "RGB")
        qimage = QImage(data, width, height, QImage.Format_RGB888)
        return QPixmap.fromImage(qimage)

    def add_row(self):
        """Add a new row to the table."""
        product_name = self.name_input.text()
        quantity = self.quantity_input.value()
        price = self.price_input.value()

        if not product_name:
            QMessageBox.warning(self, "Error", "Product Name field is empty.")
            return

        if not quantity:
            QMessageBox.warning(self, "Error", "Quantity field is empty.")
            return

        if not price:
            QMessageBox.warning(self, "Error", "Price field is empty.")
            return

        row_index = self.table.rowCount()
        self.table.insertRow(row_index)

        self.table.setItem(row_index, 0, QTableWidgetItem(product_name))
        self.table.setItem(row_index, 1, QTableWidgetItem(f"{quantity}"))
        self.table.setItem(row_index, 2, QTableWidgetItem(f"{price}"))

        # Create a smaller delete button
        delete_button = QPushButton("Delete")
        delete_button.setFixedSize(90, 25)  # Set button size (width, height)
        delete_button.clicked.connect(self.delete_row)

        # Center the text inside the button using stylesheet
        delete_button.setStyleSheet("""
            QPushButton {
                text-align: center;  /* Center the text horizontally */
                padding: 0px;        /* Remove padding to ensure text is centered */
                margin: 0px;         /* Remove margin to ensure text is centered */
            }
        """)

        # Center the button in its column
        button_widget = QWidget()
        button_layout = QHBoxLayout(button_widget)
        button_layout.addWidget(delete_button)
        button_layout.setAlignment(Qt.AlignCenter)  # Center the button
        button_layout.setContentsMargins(0, 0, 0, 0)  # Remove margins

        self.table.setCellWidget(row_index, 3, button_widget)

        # Clear input fields
        self.name_input.clear()
        self.quantity_input.setValue(0)
        self.price_input.setValue(0)

        # Update the total sum
        self.update_total_sum()

    def delete_row(self):
        """Delete a row from the table."""
        button = self.sender()
        if button:
            index = self.table.indexAt(button.pos())
            if index.isValid():
                self.table.removeRow(index.row())
                self.update_total_sum()

    def update_time(self):
        """Update the date and time label."""
        current_time = QDateTime.currentDateTime().toString("dd-MM-yyyy  hh:mm:ss")
        self.date_time_label.setText(current_time)

    def update_invoice_label(self):
        """Update the invoice label."""
        self.invoice_label.setText(f"Invoice number: {self.last_invoice_number}")

    def increment_Invoice_number(self):
        self.last_invoice_number+=1
        self.update_invoice_label()

    def reset_invoice_number(self):
        """Reset invoice number to 1 and update the UI."""
        self.last_invoice_number = 1
        self.settings.setValue("last_invoice_number", self.last_invoice_number)  # Store in QSettings
        self.invoice_label.setText(str(self.last_invoice_number))  # Update UI
        QMessageBox.information(self, "Invoice Reset", "Invoice number has been reset to 1.")


    def update_total_sum(self):
        """Calculate and update the total sum of prices in the table."""
        total_sum = 0
        for row in range(self.table.rowCount()):
            price_item = self.table.item(row, 2)
            if price_item:
                total_sum += int(price_item.text())
        self.total_sum_widget.setText(f"Total Sum: Rs {total_sum}")

    def create_policy_layout(self):
        """Create the policy notes section as a QWidget."""
        group_box = QGroupBox("Policies")  # Group box to contain policies
        layout = QVBoxLayout()
        layout.setSpacing(10)

        # Policy list widget
        self.policy_list_widget = QListWidget()

        # Add fixed policies first
        for policy in Fixed_Policy_List:
            self.add_policy_to_list(policy, fixed=True)

        # Add user-defined policies
        for policy in self.policies:
            self.add_policy_to_list(policy)

        # Input field for new policies
        self.new_policy_input = QTextEdit()
        self.new_policy_input.setPlaceholderText("Enter a new policy")
        self.new_policy_input.setFixedHeight(50)

        # Add Policy button
        add_policy_button = QPushButton("Add Policy")
        add_policy_button.clicked.connect(self.add_policy)

        # Add widgets to layout
        layout.addWidget(self.policy_list_widget)
        layout.addWidget(self.new_policy_input)
        layout.addWidget(add_policy_button)

        group_box.setLayout(layout)  # Set layout to the group box
        return group_box  # Return a QWidget (QGroupBox)




    def add_policy(self):
        """Add a new policy to the list."""
        new_policy = self.new_policy_input.toPlainText().strip()  # Use toPlainText() for QTextEdit
        if new_policy: 
            self.policies.append(new_policy)  # Add to the list
            self.add_policy_to_list(new_policy)  # Add to the UI
            self.new_policy_input.clear()  # Clear the input field

            # Save policies to settings
            self.settings.setValue("policies", self.policies)
        else:
            QMessageBox.warning(self, "Error", "Policy field is empty.")


    def delete_policy(self, item):
        """Delete a policy from the list."""
        row = self.policy_list_widget.row(item)
        if row != -1:
            policy_text = self.policies[row]  # Get text from the stored list
            del self.policies[row]  # Remove from list
            self.policy_list_widget.takeItem(row)  # Remove from UI

            # Save updated policies
            self.settings.setValue("policies", self.policies)


    def add_policy_to_list(self, policy, fixed=False):
        """Add a policy to the list widget with an optional delete button."""
        item = QListWidgetItem()
        self.policy_list_widget.addItem(item)

        # Create a QTextEdit for policy text (to enable word wrapping)
        policy_text = QTextEdit(policy)
        policy_text.setReadOnly(True)
        policy_text.setFrameStyle(QFrame.NoFrame)
        policy_text.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        policy_text.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        policy_text.setFixedHeight(50)

        # Create a widget layout
        widget = QWidget()
        widget_layout = QHBoxLayout(widget)
        widget_layout.addWidget(policy_text)

        # Only add delete button for non-fixed policies
        if not fixed:
            delete_button = QPushButton("Delete")
            delete_button.setFixedSize(60, 25)
            delete_button.clicked.connect(lambda: self.delete_policy(item))
            widget_layout.addWidget(delete_button)

        widget_layout.setContentsMargins(0, 0, 0, 0)
        item.setSizeHint(widget.sizeHint())
        self.policy_list_widget.setItemWidget(item, widget)

    """Save the entire document as a .docx file."""
    def save_as_doc(self):
        try:
            doc = Document()

            # Add logo and brand name
            doc.add_heading("Sweet Bean", level=0).alignment = 1  # Center align brand name
            logo_path = "logo.jpeg"  # Path to your logo file
            if logo_path:
                doc.add_picture(logo_path, width=Inches(1.5))  # Add logo with a width of 1.5 inches
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = 1  # Center align the logo

            # Add invoice details
            doc.add_paragraph(f"Invoice Number: {self.last_invoice_number}")
            doc.add_paragraph(f"Date and Time: {self.date_time_label.text()}")
            doc.add_paragraph(f"NTN: {self.ntn_text_label.text()}")

            # Add table
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            headers_to_save = ["Item", "Quantity", "Price"]  # Ensure headers are defined
            for i, header in enumerate(headers_to_save):
                hdr_cells[i].text = header

            for row in range(self.table.rowCount()):
                row_cells = table.add_row().cells
                for col in range(3):
                    item = self.table.item(row, col)
                    row_cells[col].text = item.text() if item else ""

            # Add total sum
            doc.add_paragraph(f"Total Sum: {self.total_sum_widget.text()}")

            # Add payment details
            doc.add_heading("Payment Details", level=2)
            doc.add_paragraph(f"Payment Method: {self.payment_method_input.text()}")
            doc.add_paragraph(f"Account Number: {self.account_number_input.text()}")
            doc.add_paragraph(f"Payment Status: {'Paid' if self.paid_radio.isChecked() else 'Pending'}")

            # Add policies
            doc.add_heading("Policies", level=2)
            
            # Add default policies
            doc.add_paragraph(Policy1)
            doc.add_paragraph(Policy2)
            
            # Add user-defined policies
            for policy in self.policies:  # Iterate through the list of policies
                doc.add_paragraph(policy)

            # Add spacing above QR codes
            doc.add_paragraph()  # Add an empty paragraph for spacing

            # Add QR codes side by side
            doc.add_heading("QR Codes", level=2)

            # Create a table with 2 columns for QR codes
            qr_table = doc.add_table(rows=1, cols=2)
            qr_table.autofit = True  # Automatically adjust column widths
            qr_cells = qr_table.rows[0].cells

            # Generate and add WhatsApp QR code
            whatsapp_qr = self.generate_qr(Whatsapp_link)
            whatsapp_qr_path = "whatsapp_qr.png"
            whatsapp_qr.save(whatsapp_qr_path)

            whatsapp_paragraph = qr_cells[0].paragraphs[0]
            run = whatsapp_paragraph.add_run()
            run.add_picture(whatsapp_qr_path, width=Inches(1.5))  # Add QR code
            whatsapp_paragraph.alignment = 1  # Center align
            whatsapp_paragraph.add_run("\nVisit our WhatsApp").bold = True  # Add label below in the same paragraph

            # Generate and add Instagram QR code
            instagram_qr = self.generate_qr(Insta_link)
            instagram_qr_path = "instagram_qr.png"
            instagram_qr.save(instagram_qr_path)

            instagram_paragraph = qr_cells[1].paragraphs[0]
            run = instagram_paragraph.add_run()
            run.add_picture(instagram_qr_path, width=Inches(1.5))  # Add QR code
            instagram_paragraph.alignment = 1  # Center align
            instagram_paragraph.add_run("\nVisit our Instagram").bold = True  # Add label below in the same paragraph

            # Save the document
            file_path, _ = QFileDialog.getSaveFileName(self, "Save Document", "", "Word Files (*.docx)")
            if file_path:
                if not file_path.endswith(".docx"):
                    file_path += ".docx"  # Ensure the file has a .docx extension
                doc.save(file_path)
                QMessageBox.information(self, "Success", "Document saved successfully!")
                self.increment_Invoice_number()  # Increment the invoice number

        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred while saving the document: {str(e)}")
        


    def closeEvent(self, event):
        """Save the last invoice number when the window is closed."""
        self.settings.setValue("last_invoice_number", self.last_invoice_number)
        event.accept()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())